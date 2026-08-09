"""
Maintains a Word file per Test Execution (TE), appending section per TC.
Stores PoC doc at: executions/<TE_KEY>/poc_report_<TE_KEY>.docx
"""
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

import config

def set_cell_background(cell, hex_color: str):
    """Sets background fill color of a docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Sets internal padding (dxa) for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """Sets subtle borders around a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def get_te_doc_path(te_key: str) -> str:
    """Returns absolute/relative path to the PoC docx file for a specific TE."""
    te_dir = os.path.join(config.EXECUTIONS_DIR, te_key)
    os.makedirs(te_dir, exist_ok=True)
    return os.path.join(te_dir, f"poc_report_{te_key}.docx")

def _get_or_create_doc(doc_path: str, te_key: str):
    if os.path.exists(doc_path):
        return Document(doc_path)
    
    doc = Document()
    
    # Compact Page Margins: 0.45 inch all around
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)

    # Document Header Table (Clean Compact Banner)
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_tbl.cell(0, 0)
    set_cell_background(cell, "1E293B")  # Dark Slate background
    set_cell_margins(cell, top=80, bottom=80, left=160, right=160)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("PROOF OF TESTING (POT) REPORT")
    r1.font.name = 'Calibri'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"Test Execution Key: {te_key}  |  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(148, 163, 184)

    # Spacing after document title
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(2)

    return doc

def rebuild_pot(te_key: str, test_cases: list) -> str:
    """Rebuilds the POT document from scratch using all test cases."""
    doc_path = get_te_doc_path(te_key)
    
    # Always create a fresh document
    doc = Document()
    
    # Compact Page Margins: 0.45 inch all around
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)

    # Document Header Table (Clean Compact Banner)
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_tbl.cell(0, 0)
    set_cell_background(cell, "1E293B")  # Dark Slate background
    set_cell_margins(cell, top=80, bottom=80, left=160, right=160)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("PROOF OF TESTING (POT) REPORT")
    r1.font.name = 'Calibri'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"Test Execution Key: {te_key}  |  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(148, 163, 184)

    # Spacing after document title
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(2)

    for tc in test_cases:
        tc_name = tc.get('key') or tc.get('name') or tc.get('tc_number') or "Unknown TC"
        status_str = (tc.get('status') or "PENDING").upper()
        summary = tc.get('summary', "")
        defect_key = tc.get('defect_key', "")
        expected_shots = tc.get('expected_shots', [])
        actual_shots = tc.get('actual_shots', [])
        
        # Get blocked TCs from submitted defect data if this is the failing TC
        blocked_tcs = ""
        if tc.get('submitted_defect'):
            blocked_tcs = tc.get('submitted_defect', {}).get('blocked_tcs', "")
        blocked_by = tc.get('blocked_by', "")

        # Status colors
        if status_str == 'PASS':
            status_bg = "DCFCE7"
            status_fg = RGBColor(22, 101, 52)
        elif status_str == 'FAIL':
            status_bg = "FEE2E2"
            status_fg = RGBColor(153, 27, 27)
        elif status_str == 'PENDING':
            status_bg = "FEF3C7"
            status_fg = RGBColor(217, 119, 6) # amber
        elif status_str == 'BLOCKED':
            status_bg = "F3E8FF"
            status_fg = RGBColor(168, 85, 247) # purple
        else: # SKIP or other
            status_bg = "F1F5F9"
            status_fg = RGBColor(100, 116, 139) # gray

        # TC Banner Table (1 row, 2 cells)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, color="CBD5E1", sz="4", val="single")

        cell_left = tbl.cell(0, 0)
        cell_right = tbl.cell(0, 1)

        cell_left.width = Inches(5.8)
        cell_right.width = Inches(1.7)

        set_cell_background(cell_left, "F8FAFC")
        set_cell_background(cell_right, status_bg)
        set_cell_margins(cell_left, top=40, bottom=40, left=100, right=100)
        set_cell_margins(cell_right, top=40, bottom=40, left=100, right=100)

        # TC Title & Info
        p_tc = cell_left.paragraphs[0]
        p_tc.paragraph_format.space_before = Pt(0)
        p_tc.paragraph_format.space_after = Pt(0)
        r_tc = p_tc.add_run(f"Test Case: {tc_name}")
        r_tc.font.name = 'Calibri'
        r_tc.font.size = Pt(10)
        r_tc.font.bold = True
        r_tc.font.color.rgb = RGBColor(15, 23, 42)

        if summary:
            p_sum = cell_left.add_paragraph()
            p_sum.paragraph_format.space_before = Pt(0)
            p_sum.paragraph_format.space_after = Pt(0)
            r_sum_lbl = p_sum.add_run("Summary: ")
            r_sum_lbl.font.bold = True
            r_sum_lbl.font.size = Pt(8.5)
            r_sum_lbl.font.color.rgb = RGBColor(71, 85, 105)
            r_sum_txt = p_sum.add_run(summary)
            r_sum_txt.font.size = Pt(8.5)
            r_sum_txt.font.color.rgb = RGBColor(51, 65, 85)

        if defect_key:
            p_def = cell_left.add_paragraph()
            p_def.paragraph_format.space_before = Pt(0)
            p_def.paragraph_format.space_after = Pt(0)
            r_def_lbl = p_def.add_run("Linked Defect: ")
            r_def_lbl.font.bold = True
            r_def_lbl.font.size = Pt(8.5)
            r_def_lbl.font.color.rgb = RGBColor(185, 28, 28)
            r_def_txt = p_def.add_run(defect_key)
            r_def_txt.font.bold = True
            r_def_txt.font.size = Pt(8.5)
            r_def_txt.font.color.rgb = RGBColor(185, 28, 28)
            
        if blocked_tcs:
            p_blk = cell_left.add_paragraph()
            p_blk.paragraph_format.space_before = Pt(0)
            p_blk.paragraph_format.space_after = Pt(0)
            r_blk_lbl = p_blk.add_run("Blocks TCs: ")
            r_blk_lbl.font.bold = True
            r_blk_lbl.font.size = Pt(8.5)
            r_blk_lbl.font.color.rgb = RGBColor(168, 85, 247)
            r_blk_txt = p_blk.add_run(blocked_tcs)
            r_blk_txt.font.size = Pt(8.5)
            r_blk_txt.font.color.rgb = RGBColor(168, 85, 247)
            
        if blocked_by:
            p_bby = cell_left.add_paragraph()
            p_bby.paragraph_format.space_before = Pt(0)
            p_bby.paragraph_format.space_after = Pt(0)
            r_bby_lbl = p_bby.add_run("Blocked By: ")
            r_bby_lbl.font.bold = True
            r_bby_lbl.font.size = Pt(8.5)
            r_bby_lbl.font.color.rgb = RGBColor(168, 85, 247)
            r_bby_txt = p_bby.add_run(blocked_by)
            r_bby_txt.font.size = Pt(8.5)
            r_bby_txt.font.color.rgb = RGBColor(168, 85, 247)

        if status_str == 'PENDING':
            p_pen = cell_left.add_paragraph()
            p_pen.paragraph_format.space_before = Pt(0)
            p_pen.paragraph_format.space_after = Pt(0)
            r_pen = p_pen.add_run("Awaiting execution")
            r_pen.font.size = Pt(8.5)
            r_pen.font.italic = True
            r_pen.font.color.rgb = RGBColor(100, 116, 139)

        # Status Pill
        p_st = cell_right.paragraphs[0]
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_st.paragraph_format.space_before = Pt(0)
        p_st.paragraph_format.space_after = Pt(0)
        r_st = p_st.add_run(f"[ {status_str} ]")
        r_st.font.name = 'Calibri'
        r_st.font.size = Pt(9.5)
        r_st.font.bold = True
        r_st.font.color.rgb = status_fg

        # Add Expected Result Screenshots
        if expected_shots:
            p_lbl_e = doc.add_paragraph()
            p_lbl_e.paragraph_format.space_before = Pt(4)
            p_lbl_e.paragraph_format.space_after = Pt(2)
            r_lbl_e = p_lbl_e.add_run("Expected Result Screenshots")
            r_lbl_e.font.bold = True
            r_lbl_e.font.size = Pt(7.5)
            r_lbl_e.font.color.rgb = RGBColor(30, 41, 59)
            
            p_img_e = doc.add_paragraph()
            p_img_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for img_path in expected_shots:
                if os.path.exists(img_path):
                    try:
                        run = p_img_e.add_run()
                        run.add_picture(img_path, height=Inches(1.5))
                        p_img_e.add_run("  ")
                    except Exception:
                        pass
        
        # Add Actual Result Screenshots
        if actual_shots:
            p_lbl_a = doc.add_paragraph()
            p_lbl_a.paragraph_format.space_before = Pt(4)
            p_lbl_a.paragraph_format.space_after = Pt(2)
            r_lbl_a = p_lbl_a.add_run("Actual Result Screenshots")
            r_lbl_a.font.bold = True
            r_lbl_a.font.size = Pt(7.5)
            r_lbl_a.font.color.rgb = RGBColor(30, 41, 59)
            
            p_img_a = doc.add_paragraph()
            p_img_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for img_path in actual_shots:
                if os.path.exists(img_path):
                    try:
                        run = p_img_a.add_run()
                        run.add_picture(img_path, height=Inches(1.5))
                        p_img_a.add_run("  ")
                    except Exception:
                        pass

        # Divider spacing after TC entry
        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_before = Pt(4)
        p_div.paragraph_format.space_after = Pt(4)

    doc.save(doc_path)
    return doc_path

def append_tc_pot(
    tc_number: str, 
    status: str, 
    te_key: str = None, 
    summary: str = "", 
    expected_shots: list[str] = None, 
    actual_shots: list[str] = None,
    defect_key: str = None,
    blocked_tcs: str = None
):
    """
    Appends a compact TC result to the POT doc, enabling 3 TCs per page.
    """
    if not te_key:
        te_key = config.DEFAULT_TE_KEY

    expected_shots = [p for p in (expected_shots or []) if os.path.exists(p)]
    actual_shots = [p for p in (actual_shots or []) if os.path.exists(p)]

    doc_path = get_te_doc_path(te_key)
    doc = _get_or_create_doc(doc_path, te_key)

    status_str = (status or "PASS").upper()
    
    # Status colors
    if status_str == 'PASS':
        status_bg = "DCFCE7"
        status_fg = RGBColor(22, 101, 52)
    elif status_str == 'FAIL':
        status_bg = "FEE2E2"
        status_fg = RGBColor(153, 27, 27)
    elif status_str == 'BLOCKED':
        status_bg = "F3E8FF"
        status_fg = RGBColor(168, 85, 247)
    else:
        status_bg = "FEF3C7"
        status_fg = RGBColor(146, 64, 14)

    # TC Banner Table (1 row, 2 cells)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="CBD5E1", sz="4", val="single")

    cell_left = tbl.cell(0, 0)
    cell_right = tbl.cell(0, 1)

    cell_left.width = Inches(5.8)
    cell_right.width = Inches(1.7)

    set_cell_background(cell_left, "F8FAFC")
    set_cell_background(cell_right, status_bg)
    set_cell_margins(cell_left, top=40, bottom=40, left=100, right=100)
    set_cell_margins(cell_right, top=40, bottom=40, left=100, right=100)

    # TC Title & Info
    p_tc = cell_left.paragraphs[0]
    p_tc.paragraph_format.space_before = Pt(0)
    p_tc.paragraph_format.space_after = Pt(0)
    r_tc = p_tc.add_run(f"Test Case: {tc_number}")
    r_tc.font.name = 'Calibri'
    r_tc.font.size = Pt(10)
    r_tc.font.bold = True
    r_tc.font.color.rgb = RGBColor(15, 23, 42)

    if summary:
        p_sum = cell_left.add_paragraph()
        p_sum.paragraph_format.space_before = Pt(0)
        p_sum.paragraph_format.space_after = Pt(0)
        r_sum_lbl = p_sum.add_run("Summary: ")
        r_sum_lbl.font.bold = True
        r_sum_lbl.font.size = Pt(8.5)
        r_sum_lbl.font.color.rgb = RGBColor(71, 85, 105)
        r_sum_txt = p_sum.add_run(summary)
        r_sum_txt.font.size = Pt(8.5)
        r_sum_txt.font.color.rgb = RGBColor(51, 65, 85)

    if defect_key:
        p_def = cell_left.add_paragraph()
        p_def.paragraph_format.space_before = Pt(0)
        p_def.paragraph_format.space_after = Pt(0)
        r_def_lbl = p_def.add_run("Linked Defect: ")
        r_def_lbl.font.bold = True
        r_def_lbl.font.size = Pt(8.5)
        r_def_lbl.font.color.rgb = RGBColor(185, 28, 28)
        r_def_txt = p_def.add_run(defect_key)
        r_def_txt.font.bold = True
        r_def_txt.font.size = Pt(8.5)
        r_def_txt.font.color.rgb = RGBColor(185, 28, 28)

    if blocked_tcs:
        p_blk = cell_left.add_paragraph()
        p_blk.paragraph_format.space_before = Pt(0)
        p_blk.paragraph_format.space_after = Pt(0)
        r_blk_lbl = p_blk.add_run("Blocks TCs: ")
        r_blk_lbl.font.bold = True
        r_blk_lbl.font.size = Pt(8.5)
        r_blk_lbl.font.color.rgb = RGBColor(168, 85, 247)
        r_blk_txt = p_blk.add_run(blocked_tcs)
        r_blk_txt.font.size = Pt(8.5)
        r_blk_txt.font.color.rgb = RGBColor(168, 85, 247)

    # Status Pill
    p_st = cell_right.paragraphs[0]
    p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_st.paragraph_format.space_before = Pt(0)
    p_st.paragraph_format.space_after = Pt(0)
    r_st = p_st.add_run(f"[ {status_str} ]")
    r_st.font.name = 'Calibri'
    r_st.font.size = Pt(9.5)
    r_st.font.bold = True
    r_st.font.color.rgb = status_fg

    # Image Section — Side by side compact layout
    has_exp = len(expected_shots) > 0
    has_act = len(actual_shots) > 0

    MAX_IMG_HEIGHT = Inches(1.0)

    if has_exp and has_act:
        # 2-column side-by-side table comparing Expected vs Actual
        img_tbl = doc.add_table(rows=2, cols=2)
        img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(img_tbl, color="E2E8F0", sz="4", val="single")

        # Headers
        c_hdr_exp = img_tbl.cell(0, 0)
        c_hdr_act = img_tbl.cell(0, 1)

        c_hdr_exp.width = Inches(3.75)
        c_hdr_act.width = Inches(3.75)

        set_cell_background(c_hdr_exp, "F1F5F9")
        set_cell_background(c_hdr_act, "F1F5F9")
        set_cell_margins(c_hdr_exp, top=30, bottom=30, left=60, right=60)
        set_cell_margins(c_hdr_act, top=30, bottom=30, left=60, right=60)

        p_h1 = c_hdr_exp.paragraphs[0]
        p_h1.paragraph_format.space_before = Pt(0)
        p_h1.paragraph_format.space_after = Pt(0)
        r_h1 = p_h1.add_run("Expected Result (Figma / Design)")
        r_h1.font.bold = True
        r_h1.font.size = Pt(7.5)
        r_h1.font.color.rgb = RGBColor(30, 41, 59)

        p_h2 = c_hdr_act.paragraphs[0]
        p_h2.paragraph_format.space_before = Pt(0)
        p_h2.paragraph_format.space_after = Pt(0)
        r_h2 = p_h2.add_run("Actual Result (Test Evidence)")
        r_h2.font.bold = True
        r_h2.font.size = Pt(7.5)
        r_h2.font.color.rgb = RGBColor(30, 41, 59)

        # Image Cells
        c_img_exp = img_tbl.cell(1, 0)
        c_img_act = img_tbl.cell(1, 1)
        c_img_exp.width = Inches(3.75)
        c_img_act.width = Inches(3.75)
        set_cell_margins(c_img_exp, top=40, bottom=40, left=40, right=40)
        set_cell_margins(c_img_act, top=40, bottom=40, left=40, right=40)

        # Populate Expected images (side-by-side in single paragraph line)
        p_exp = c_img_exp.paragraphs[0]
        p_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_exp.paragraph_format.space_before = Pt(0)
        p_exp.paragraph_format.space_after = Pt(0)

        for idx, img_path in enumerate(expected_shots):
            if idx > 0:
                p_exp.add_run("  ")
            try:
                run = p_exp.add_run()
                run.add_picture(img_path, height=MAX_IMG_HEIGHT)
            except Exception as e:
                p_exp.add_run(f"[Image load error: {e}]")

        # Populate Actual images (side-by-side in single paragraph line)
        p_act = c_img_act.paragraphs[0]
        p_act.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_act.paragraph_format.space_before = Pt(0)
        p_act.paragraph_format.space_after = Pt(0)

        for idx, img_path in enumerate(actual_shots):
            if idx > 0:
                p_act.add_run("  ")
            try:
                run = p_act.add_run()
                run.add_picture(img_path, height=MAX_IMG_HEIGHT)
            except Exception as e:
                p_act.add_run(f"[Image load error: {e}]")

    elif has_exp or has_act:
        shots = expected_shots if has_exp else actual_shots
        title_text = "Expected Result (Figma / Design)" if has_exp else "Actual Result (Test Evidence)"

        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(2)
        p_lbl.paragraph_format.space_after = Pt(2)
        r_lbl = p_lbl.add_run(title_text)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.color.rgb = RGBColor(30, 41, 59)

        # Lay images out in a grid table
        num_cols = min(4, len(shots))
        num_rows = (len(shots) + num_cols - 1) // num_cols if num_cols > 0 else 1
        grid_tbl = doc.add_table(rows=num_rows, cols=max(1, num_cols))
        grid_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(grid_tbl, color="E2E8F0", sz="4", val="single")

        for idx, img_path in enumerate(shots):
            r_i = idx // num_cols
            c_i = idx % num_cols
            c = grid_tbl.cell(r_i, c_i)
            c.width = Inches(7.5 / num_cols)
            set_cell_margins(c, top=40, bottom=40, left=40, right=40)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            try:
                run = p.add_run()
                run.add_picture(img_path, height=MAX_IMG_HEIGHT)
            except Exception as e:
                p.add_run(f"[Image error: {e}]")

    # Divider spacing after TC entry
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(2)
    p_div.paragraph_format.space_after = Pt(2)

    doc.save(doc_path)
    print(f"  -> {tc_number} POT entry saved to {doc_path}")
    return doc_path
    return doc_path

def append_pass(tc_key: str, screenshot_paths: list[str] = None, te_key: str = None, summary: str = ""):
    """Helper for pass workflow in CLI/main.py."""
    return append_tc_pot(
        tc_number=tc_key,
        status="PASS",
        te_key=te_key,
        summary=summary or "Test Case Passed",
        actual_shots=screenshot_paths
    )

